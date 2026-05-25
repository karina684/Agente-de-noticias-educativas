import requests
import json
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

ANTHROPIC_API_KEY = os.environ['ANTHROPIC_API_KEY']
EMAIL_FROM = os.environ['EMAIL_FROM']
EMAIL_PASSWORD = os.environ['EMAIL_PASSWORD']
EMAIL_TO = 'karina@borlenghi.com'

hoy = datetime.now()
es_lunes = hoy.weekday() == 0
if es_lunes:
    periodo = "del fin de semana"
    dias = "los últimos 3 días"
else:
    periodo = "de las últimas 24 horas"
    dias = "el último día"

def buscar_noticias():
    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "Content-Type": "application/json",
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "anthropic-beta": "search-grounding-2025-02-19"
        },
        json={
            "model": "claude-opus-4-5",
            "max_tokens": 4000,
            "tools": [{"type": "web_search_20250305", "name": "web_search"}],
            "messages": [{
                "role": "user",
                "content": f"""Buscá noticias educativas argentinas de {dias} en estos medios: Clarín, La Nación, Infobae, Página/12, y también resoluciones educativas del Boletín Oficial Nacional (boletinoficial.gob.ar) y del Boletín Oficial de CABA (boletinoficial.buenosaires.gob.ar).

Temas: escuelas, universidades, docentes, estudiantes, políticas educativas, normativas de educación.

Devolvé SOLO un JSON válido con este formato exacto, sin texto adicional:
{{
  "noticias": [
    {{
      "titulo": "título de la noticia",
      "url": "url completa",
      "fuente": "nombre del diario o Boletín Oficial Nacional/CABA",
      "numero_norma": "número de resolución si es boletín oficial, sino vacío",
      "fecha": "DD/MM/YYYY",
      "hora": "HH:MM",
      "resumen": "resumen de 2 a 3 líneas"
    }}
  ]
}}"""
            }]
        }
    )
    
    data = response.json()
    texto = ""
    for block in data.get("content", []):
        if block.get("type") == "text":
            texto += block.get("text", "")
    
    texto = texto.strip()
    if "```json" in texto:
        texto = texto.split("```json")[1].split("```")[0].strip()
    elif "```" in texto:
        texto = texto.split("```")[1].split("```")[0].strip()
    
    return json.loads(texto)

def generar_word(noticias):
    doc = Document()
    
    # Título
    titulo = doc.add_heading('Resumen de Noticias Educativas', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitulo = doc.add_paragraph(f'Fecha: {hoy.strftime("%d/%m/%Y")} — Noticias {periodo}')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    for n in noticias:
        # Título de la noticia
        p_titulo = doc.add_paragraph()
        run = p_titulo.add_run(n.get('titulo', ''))
        run.bold = True
        run.font.size = Pt(12)
        
        # Fuente o número de norma
        fuente = n.get('fuente', '')
        norma = n.get('numero_norma', '')
        if norma:
            p_fuente = doc.add_paragraph()
            p_fuente.add_run(f'Norma: {norma} | {fuente}').italic = True
        else:
            p_fuente = doc.add_paragraph()
            p_fuente.add_run(f'Fuente: {fuente}').italic = True
        
        # Fecha y hora
        fecha = n.get('fecha', '')
        hora = n.get('hora', '')
        p_fecha = doc.add_paragraph()
        p_fecha.add_run(f'Fecha: {fecha} {hora}').italic = True
        
        # URL
        url = n.get('url', '')
        if url:
            p_url = doc.add_paragraph()
            p_url.add_run(f'Link: {url}').italic = True
        
        # Resumen
        doc.add_paragraph(n.get('resumen', ''))
        
        # Separador
        doc.add_paragraph('─' * 60)
    
    archivo = 'noticias_educativas.docx'
    doc.save(archivo)
    return archivo

def enviar_email(archivo):
    msg = MIMEMultipart()
    msg['From'] = EMAIL_FROM
    msg['To'] = EMAIL_TO
    msg['Subject'] = f'📚 Noticias Educativas — {hoy.strftime("%d/%m/%Y")}'
    
    cuerpo = f'Adjunto el resumen de noticias educativas {periodo}.'
    msg.attach(MIMEText(cuerpo, 'plain'))
    
    with open(archivo, 'rb') as f:
        part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename={archivo}')
    msg.attach(part)
    
    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.starttls()
        server.login(EMAIL_FROM, EMAIL_PASSWORD)
        server.send_message(msg)
    
    print(f'✅ Email enviado a {EMAIL_TO}')

# Ejecutar
print('Buscando noticias...')
resultado = buscar_noticias()
noticias = resultado.get('noticias', [])
print(f'Encontradas: {len(noticias)} noticias')

archivo = generar_word(noticias)
print(f'Word generado: {archivo}')

enviar_email(archivo)
