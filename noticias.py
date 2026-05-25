import requests
import json
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

ANTHROPIC_API_KEY = os.environ['ANTHROPIC_API_KEY']
EMAIL_FROM = os.environ['EMAIL_FROM']
EMAIL_PASSWORD = os.environ['EMAIL_PASSWORD']
EMAIL_TO = 'karina@borlenghi.com'

hoy = datetime.now()
es_lunes = hoy.weekday() == 0
periodo = "del fin de semana" if es_lunes else "de las últimas 24 horas"
dias = "los últimos 3 días" if es_lunes else "el último día"

def buscar_noticias():
    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "Content-Type": "application/json",
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01"
        },
        json={
            "model": "claude-opus-4-5",
            "max_tokens": 4000,
            "tools": [{"type": "web_search_20250305", "name": "web_search"}],
            "messages": [{
                "role": "user",
                "content": f"""Buscá noticias educativas argentinas de {dias} en: Clarín, La Nación, Infobae, Página/12, Boletín Oficial Nacional (boletinoficial.gob.ar) y Boletín Oficial CABA (boletinoficial.buenosaires.gob.ar). Temas: escuelas, universidades, docentes, estudiantes, políticas educativas, normativas.

Devolvé SOLO JSON válido sin texto adicional:
{{"noticias": [{{"titulo": "título", "url": "url completa", "fuente": "nombre del medio", "numero_norma": "número si es boletín oficial sino vacío", "fecha": "DD/MM/YYYY", "hora": "HH:MM", "resumen": "2 a 3 líneas"}}]}}"""
            }]
        }
    )
    data = response.json()
    texto = ""
    for block in data.get("content", []):
        if block.get("type") == "text":
            texto += block.get("text", "")
    texto = texto.strip()
    if "```json" in texto:
        texto = texto.split("```json")[1].split("```")[0].strip()
    elif "```" in texto:
        texto = texto.split("```")[1].split("```")[0].strip()
    return json.loads(texto)

def generar_word(noticias):
    doc = Document()
    titulo = doc.add_heading('Resumen de Noticias Educativas', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Fecha: {hoy.strftime("%d/%m/%Y")} — Noticias {periodo}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    for n in noticias:
        p = doc.add_paragraph()
        r = p.add_run(n.get('titulo', ''))
        r.bold = True
        r.font.size = Pt(12)
        norma = n.get('numero_norma', '')
        fuente = n.get('fuente', '')
        p2 = doc.add_paragraph()
        p2.add_run(f'{"Norma: " + norma + " | " if norma else ""}Fuente: {fuente}').italic = True
        p3 = doc.add_paragraph()
        p3.add_run(f'Fecha: {n.get("fecha","")} {n.get("hora","")}').italic = True
        url = n.get('url', '')
        if url:
            p4 = doc.add_paragraph()
            p4.add_run(f'Link: {url}').italic = True
        doc.add_paragraph(n.get('resumen', ''))
        doc.add_paragraph('─' * 60)
    archivo = 'noticias_educativas.docx'
    doc.save(archivo)
    return archivo

def enviar_email(archivo):
    msg = MIMEMultipart()
    msg['From'] = EMAIL_FROM
    msg['To'] = EMAIL_TO
    msg['Subject'] = f'Noticias Educativas — {hoy.strftime("%d/%m/%Y")}'
    msg.attach(MIMEText(f'Adjunto el resumen de noticias educativas {periodo}.', 'plain'))
    with open(archivo, 'rb') as f:
        part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename={archivo}')
    msg.attach(part)
    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.starttls()
        server.login(EMAIL_FROM, EMAIL_PASSWORD)
        server.send_message(msg)
    print(f'Email enviado a {EMAIL_TO}')

print('Buscando noticias...')
resultado = buscar_noticias()
noticias = resultado.get('noticias', [])
print(f'Encontradas: {len(noticias)} noticias')
archivo = generar_word(noticias)
print(f'Word generado: {archivo}')
enviar_email(archivo)
